# ==========================
# 🔧 Standard Library Imports
# ==========================
import os
import re
from datetime import datetime, timedelta
from io import BytesIO
import json
import base64
from math import ceil

import requests
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from sqlalchemy import or_, and_, text

# ==========================
# 🧩 Third-Party Imports
# ==========================
from flask import (Flask, Response, flash, redirect, render_template, request,
                   send_from_directory, session, url_for, jsonify)
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS

# Configure Flask to serve static files from the React build directory
FRONTEND_DIST_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'frontend', 'dist')

app = Flask(__name__, static_folder=os.path.join(FRONTEND_DIST_DIR, 'assets'), template_folder='templates')
CORS(app)

load_dotenv()

app.secret_key = os.getenv("SECRET_KEY", "dev_key")
ADMIN_ID = os.getenv("ADMIN_ID")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD")

# Database Config
database_url = os.getenv("DATABASE_URL")
if database_url and database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://", 1)

if not database_url:
    raise ValueError("DATABASE_URL environment variable is not set!")

app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)

# -------------------------------
# MODELS
# -------------------------------

post_tags = db.Table('post_tags',
    db.Column('post_id', db.Integer, db.ForeignKey('post.id'), primary_key=True),
    db.Column('tag_id', db.Integer, db.ForeignKey('tag.id'), primary_key=True)
)

class Tag(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False, unique=True)

class Category(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False, unique=True)
    description = db.Column(db.String(255), nullable=True)
    posts = db.relationship('Post', backref='category', lazy=True)

class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(150), nullable=False)
    content = db.Column(db.Text, nullable=False)
    date_posted = db.Column(db.DateTime, default=datetime.utcnow)
    views = db.Column(db.Integer, default=0)
    category_id = db.Column(db.Integer, db.ForeignKey('category.id'), nullable=True)
    tags = db.relationship('Tag', secondary=post_tags, lazy='subquery', backref=db.backref('posts', lazy=True))

with app.app_context():
    db.create_all()
    if not Category.query.first():
        default_cat = Category(name="General", description="General topics")
        db.session.add(default_cat)
        db.session.commit()

# -------------------------------
# ADMIN HELPERS & ROUTES
# -------------------------------
def is_admin():
    return session.get("is_admin")

@app.route("/zytez-login", methods=["GET", "POST"])
def zytez_login():
    if request.method == "POST":
        user_id = request.form.get("user_id")
        password = request.form.get("password")
        if user_id == ADMIN_ID and password == ADMIN_PASSWORD:
            session["is_admin"] = True
            session.permanent = True
            app.permanent_session_lifetime = timedelta(minutes=30)
            flash("✅ Giriş başarılı!", "success")
            return redirect(url_for("zytez_dashboard"))
        else:
            flash("❌ Hatalı ID veya şifre!", "error")
            return redirect(url_for("zytez_login"))
    return render_template("zytez_login.html")

@app.route("/zytez")
def zytez_dashboard():
    if not is_admin(): return redirect(url_for("zytez_login"))
    return render_template("zytez_dashboard.html")

# --- Post Management ---
@app.route("/zytez/posts")
def manage_posts():
    if not is_admin(): return redirect(url_for("zytez_login"))
    posts = Post.query.order_by(Post.date_posted.desc()).all()
    return render_template("manage_posts.html", posts=posts)

@app.route("/zytez/posts/add", methods=["GET", "POST"])
def add_post():
    if not is_admin(): return redirect(url_for("zytez_login"))
    if request.method == "POST":
        try:
            title = request.form.get('title')
            content = request.form.get('content')
            category_id = request.form.get('category_id')
            
            if not title or not content or not category_id:
                flash("⚠️ Başlık, içerik ve kategori zorunludur!", "error")
                return redirect(url_for("add_post"))

            new_post = Post(
                title=title, 
                content=content, 
                category_id=int(category_id)
            )
            
            tag_ids = request.form.getlist('tags')
            for tag_id in tag_ids:
                tag = Tag.query.get(int(tag_id))
                if tag:
                    new_post.tags.append(tag)
            
            db.session.add(new_post)
            db.session.commit()
            flash("✅ Yeni yazı eklendi!", "success")
            return redirect(url_for("manage_posts"))
        except Exception as e:
            db.session.rollback()
            print(f"Error adding post: {e}")
            flash(f"❌ Hata oluştu: {str(e)}", "error")
            return redirect(url_for("add_post"))
    
    categories = Category.query.all()
    tags = Tag.query.all()
    return render_template("edit_post.html", post=None, categories=categories, tags=tags)

@app.route("/zytez/posts/edit/<int:post_id>", methods=["GET", "POST"])
def edit_post(post_id):
    if not is_admin(): return redirect(url_for("zytez_login"))
    post = Post.query.get_or_404(post_id)
    if request.method == "POST":
        try:
            post.title = request.form.get('title')
            post.content = request.form.get('content')
            post.category_id = int(request.form.get('category_id'))
            
            post.tags.clear()
            tag_ids = request.form.getlist('tags')
            for tag_id in tag_ids:
                tag = Tag.query.get(int(tag_id))
                if tag:
                    post.tags.append(tag)
            
            db.session.commit()
            flash("✅ Yazı güncellendi!", "success")
            return redirect(url_for("manage_posts"))
        except Exception as e:
            db.session.rollback()
            print(f"Error editing post: {e}")
            flash(f"❌ Hata oluştu: {str(e)}", "error")
            return redirect(url_for("edit_post", post_id=post_id))

    categories = Category.query.all()
    tags = Tag.query.all()
    return render_template("edit_post.html", post=post, categories=categories, tags=tags)

@app.route("/zytez/posts/delete/<int:post_id>", methods=["POST"])
def delete_post(post_id):
    if not is_admin(): return redirect(url_for("zytez_login"))
    post = Post.query.get_or_404(post_id)
    db.session.delete(post)
    db.session.commit()
    flash("🗑️ Yazı silindi.", "success")
    return redirect(url_for("manage_posts"))

# --- Category Management ---
@app.route("/zytez/categories")
def manage_categories():
    if not is_admin(): return redirect(url_for("zytez_login"))
    categories = Category.query.all()
    return render_template("manage_categories.html", categories=categories)

@app.route("/zytez/categories/add", methods=["POST"])
def add_category():
    if not is_admin(): return redirect(url_for("zytez_login"))
    name = request.form.get("name")
    if name and not Category.query.filter_by(name=name).first():
        new_cat = Category(name=name, description=request.form.get("description"))
        db.session.add(new_cat)
        db.session.commit()
        flash("✅ Yeni kategori eklendi!", "success")
    else:
        flash("⚠️ Kategori adı boş veya zaten var!", "error")
    return redirect(url_for("manage_categories"))

@app.route("/zytez/categories/delete/<int:cat_id>", methods=["POST"])
def delete_category(cat_id):
    if not is_admin(): return redirect(url_for("zytez_login"))
    category = Category.query.get_or_404(cat_id)
    for post in category.posts:
        post.category_id = None
    db.session.delete(category)
    db.session.commit()
    flash("🗑️ Kategori silindi.", "success")
    return redirect(url_for("manage_categories"))

# --- Tag Management ---
@app.route("/zytez/tags")
def manage_tags():
    if not is_admin(): return redirect(url_for("zytez_login"))
    tags = Tag.query.all()
    return render_template("manage_tags.html", tags=tags)

@app.route("/zytez/tags/add", methods=["POST"])
def add_tag():
    if not is_admin(): return redirect(url_for("zytez_login"))
    name = request.form.get("name")
    if name and not Tag.query.filter_by(name=name).first():
        new_tag = Tag(name=name)
        db.session.add(new_tag)
        db.session.commit()
        flash("✅ Yeni etiket eklendi!", "success")
    else:
        flash("⚠️ Etiket adı boş veya zaten var!", "error")
    return redirect(url_for("manage_tags"))

@app.route("/zytez/tags/delete/<int:tag_id>", methods=["POST"])
def delete_tag(tag_id):
    if not is_admin(): return redirect(url_for("zytez_login"))
    tag = Tag.query.get_or_404(tag_id)
    db.session.delete(tag)
    db.session.commit()
    flash("🗑️ Etiket silindi.", "success")
    return redirect(url_for("manage_tags"))

# --- DB Management ---
@app.route("/zytez/database")
def manage_database():
    if not is_admin(): return redirect(url_for("zytez_login"))
    return render_template("manage_db.html")

# -------------------------------
# API ROUTES
# -------------------------------
@app.route("/api/posts")
def api_posts():
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 10, type=int)
        category_slug = request.args.get('category')
        search_query = request.args.get('q')
        
        query = Post.query
        
        if category_slug:
            clean_name = category_slug.replace('-', ' ')
            category = Category.query.filter(Category.name.ilike(clean_name)).first()
            if category:
                query = query.filter_by(category_id=category.id)
            else:
                return jsonify({'posts': [], 'total': 0, 'pages': 0})
        
        if search_query:
            candidates = query.filter(
                (Post.title.ilike(f"%{search_query}%")) | 
                (Post.content.ilike(f"%{search_query}%"))
            ).order_by(Post.date_posted.desc()).all()
            
            final_posts = []
            for p in candidates:
                if search_query.lower() in p.title.lower():
                    final_posts.append(p)
                    continue
                
                soup = BeautifulSoup(p.content, 'html.parser')
                text_content = soup.get_text()
                
                if re.search(r'\b' + re.escape(search_query) + r'\b', text_content, re.IGNORECASE):
                    final_posts.append(p)
            
            total = len(final_posts)
            start = (page - 1) * per_page
            end = start + per_page
            paginated_posts = final_posts[start:end]
            total_pages = ceil(total / per_page)
            
            posts = paginated_posts
        else:
            pagination = query.order_by(Post.date_posted.desc()).paginate(page=page, per_page=per_page, error_out=False)
            posts = pagination.items
            total = pagination.total
            total_pages = pagination.pages
        
        data = []
        for p in posts:
            soup = BeautifulSoup(p.content, 'html.parser')
            text = soup.get_text()
            excerpt = text[:200] + "..." if len(text) > 200 else text
            
            cat_name = p.category.name if p.category else "Uncategorized"
            tag_names = [t.name for t in p.tags]
            
            data.append({
                "id": str(p.id),
                "title": p.title,
                "excerpt": excerpt,
                "date": p.date_posted.strftime("%B %d, %Y"),
                "author": "Yunus Tez",
                "tags": tag_names,
                "category": cat_name
            })
            
        return jsonify({
            'posts': data,
            'total': total,
            'pages': total_pages,
            'current_page': page
        })
    except Exception as e:
        print(f"Error fetching posts: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/posts/<int:id>")
def api_post_detail(id):
    try:
        p = Post.query.get_or_404(id)
        p.views += 1
        db.session.commit()
        
        cat_name = p.category.name if p.category else "Uncategorized"
        tag_names = [t.name for t in p.tags]
        
        return jsonify({
            "id": str(p.id),
            "title": p.title,
            "content": p.content,
            "date": p.date_posted.strftime("%B %d, %Y"),
            "author": "Yunus Tez",
            "tags": tag_names,
            "category": cat_name
        })
    except Exception as e:
        print(f"Error fetching post detail: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/categories")
def api_categories():
    try:
        categories = Category.query.all()
        data = []
        for c in categories:
            post_count = Post.query.filter_by(category_id=c.id).count()
            data.append({
                "id": str(c.id),
                "name": c.name,
                "description": c.description or "",
                "count": post_count,
                "color": "#0A6ED1"
            })
        return jsonify(data)
    except Exception as e:
        print(f"Error fetching categories: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/tags")
def api_tags():
    try:
        tags = Tag.query.all()
        data = [t.name for t in tags]
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/", defaults={'path': ''})
@app.route('/<path:path>')
def serve(path):
    if path != "" and os.path.exists(os.path.join(FRONTEND_DIST_DIR, path)):
        return send_from_directory(FRONTEND_DIST_DIR, path)
    else:
        if not (path.startswith("api") or path.startswith("zytez")):
            return send_from_directory(FRONTEND_DIST_DIR, 'index.html')
        else:
            return "Not Found", 404

# -------------------------------
# UTILS (Backup/Restore)
# -------------------------------
@app.route('/backup/docx')
def backup_docx():
    if not is_admin(): return redirect(url_for("zytez_login"))
    posts = Post.query.order_by(Post.date_posted.desc()).all()
    doc = Document()
    doc.add_heading('📚 Blog Yedek Raporu', 0)

    for post in posts:
        doc.add_heading(post.title, level=1)
        doc.add_paragraph(f"Tarih: {post.date_posted.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph('')

        soup = BeautifulSoup(post.content, 'html.parser')
        for element in soup.find_all(['p', 'img']):
            if element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            elif element.name == 'img':
                src = element.get('src')
                if not src: continue
                if src.startswith('data:image'):
                    try:
                        header, data = src.split(',', 1)
                        data = data.replace('\n', '')
                        image_data = base64.b64decode(data)
                        doc.add_picture(BytesIO(image_data), width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Görsel çözülemedi: {e}]")
                elif src.startswith('http'):
                    try:
                        img_data = requests.get(src, timeout=10).content
                        doc.add_picture(BytesIO(img_data), width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Görsel indirilemedi: {src}]")
        doc.add_page_break()

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return Response(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=blog_backup.docx"}
    )

@app.route("/backup/json")
def backup_json():
    if not is_admin(): return redirect(url_for("zytez_login"))
    posts = Post.query.all()
    categories = Category.query.all()
    tags = Tag.query.all()
    
    data = {
        "categories": [
            {"id": c.id, "name": c.name, "description": c.description}
            for c in categories
        ],
        "tags": [
            {"id": t.id, "name": t.name} for t in tags
        ],
        "posts": [
            {
                "id": p.id,
                "title": p.title,
                "content": p.content,
                "views": p.get("views", 0) if hasattr(p, "views") else 0,
                "date_posted": p.date_posted.strftime("%Y-%m-%d"),
                "category_id": p.category_id,
                "tags": [t.name for t in p.tags]
            }
            for p in posts
        ]
    }
    return Response(
        json.dumps(data, ensure_ascii=False, indent=4),
        mimetype="application/json",
        headers={"Content-Disposition": "attachment;filename=db_dump.json"}
    )

def fix_sequences():
    """PostgreSQL sequence'larını tablodaki max ID'ye göre düzeltir."""
    try:
        # PostgreSQL'e özgü sequence düzeltme komutları
        db.session.execute(text("SELECT setval(pg_get_serial_sequence('post', 'id'), coalesce(max(id),0) + 1, false) FROM post;"))
        db.session.execute(text("SELECT setval(pg_get_serial_sequence('category', 'id'), coalesce(max(id),0) + 1, false) FROM category;"))
        db.session.execute(text("SELECT setval(pg_get_serial_sequence('tag', 'id'), coalesce(max(id),0) + 1, false) FROM tag;"))
        db.session.commit()
        print("✅ Sequences fixed successfully.")
    except Exception as e:
        print(f"⚠️ Could not fix sequences (might be SQLite): {e}")
        db.session.rollback()

@app.route("/restore/json", methods=["POST"])
def restore_json():
    if not is_admin(): return redirect(url_for("zytez_login"))
    file = request.files.get("dumpfile")
    if not file:
        flash("❌ Dosya bulunamadı", "error")
        return redirect(url_for("manage_database"))
    try:
        data = json.load(file)
    except:
        flash("❌ JSON formatı hatalı!", "error")
        return redirect(url_for("manage_database"))

    # Clear existing data
    db.session.execute(post_tags.delete())
    Post.query.delete()
    Category.query.delete()
    Tag.query.delete()
    db.session.commit()

    # Restore categories
    for c in data.get("categories", []):
        new_cat = Category(id=c.get("id"), name=c.get("name"), description=c.get("description"))
        db.session.add(new_cat)
    
    # Restore tags
    for t in data.get("tags", []):
        new_tag = Tag(id=t.get("id"), name=t.get("name"))
        db.session.add(new_tag)
        
    db.session.commit()

    # Restore posts
    for p in data.get("posts", []):
        new_post = Post(
            id=p.get("id"),
            title=p.get("title"),
            content=p.get("content"),
            views=p.get("views", 0),
            date_posted = datetime.strptime(p.get("date_posted"), "%Y-%m-%d"),
            category_id=p.get("category_id")
        )
        
        if "tags" in p:
            for t_name in p["tags"]:
                tag = Tag.query.filter_by(name=t_name).first()
                if tag:
                    new_post.tags.append(tag)
                    
        db.session.add(new_post)
    db.session.commit()
    
    # Fix sequences after restore
    fix_sequences()
    
    flash("✅ Veritabanı başarıyla geri yüklendi!", "success")
    return redirect(url_for("manage_database"))

@app.route("/zytez/fix-sequences")
def manual_fix_sequences():
    if not is_admin(): return redirect(url_for("zytez_login"))
    fix_sequences()
    flash("✅ Veritabanı sayaçları düzeltildi.", "success")
    return redirect(url_for("manage_database"))

@app.route("/sitemap.xml", methods=["GET"])
def sitemap():
    try:
        posts = Post.query.order_by(Post.date_posted.desc()).all()
        base_url = "https://yunustez.com.tr"
        lastmod = posts[0].date_posted if posts else datetime.utcnow()
        xml = render_template("sitemap.xml", posts=posts, base_url=base_url, lastmod=lastmod)
        return Response(xml, mimetype="application/xml")
    except Exception:
        return "Sitemap not available", 404

@app.route("/logout")
def logout():
    session.pop("is_admin", None)
    flash("👋 Oturum kapatıldı.", "info")
    return redirect("/")

if __name__ == "__main__":
    app.run(debug=True)